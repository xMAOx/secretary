# -*- coding: UTF-8 -*-
from django.shortcuts import render

from django.shortcuts import render_to_response, redirect
from django.http import HttpResponse
from django.template import RequestContext
from models import Diary, Month
from forms import DiaryForm
from django.core.exceptions import ObjectDoesNotExist
from django.utils import timezone
from django.utils.timezone import localtime

from django.shortcuts import render_to_response, redirect
from django.http import HttpResponse
from django.template import RequestContext
from models import Diary
import StringIO
from docx import *
from docx.shared import Inches

# 瀏覽日誌
def diary(request):
        diaries = Diary.objects.all().order_by("-id")
        return render_to_response('diary.html', {'diaries': diaries}, context_instance=RequestContext(request))

def diary_add(request):
        if request.method == 'POST':
                form = DiaryForm(request.POST)
                if form.is_valid():
                        form.save()
                        year = localtime(timezone.now()).year
                        month =  localtime(timezone.now()).month
                        try:
                                themonth = Month.objects.get(date=year*100+month)
                        except ObjectDoesNotExist:
                                themonth = Month(date=year*100+month)
                                themonth.save()
                        return redirect("/diary")
        else:
                form = DiaryForm()
        return render_to_response('form.html',{'form': form}, context_instance=RequestContext(request))
def home(request):
        months = Month.objects.all().order_by("-id")
        return render_to_response('home.html', {'months': months}, context_instance=RequestContext(request))
def diary_word(request, month):
        document = Document()
        docx_title="Diary-"+str(timezone.localtime(timezone.now()).date())+".docx"

        time_year = int(month)/100
        time_month = int(month)%100
        diaries = Diary.objects.filter(time__year=time_year, time__month=time_month).order_by("-id")
        paragraph = document.add_paragraph(u'我的日誌：'+month)
        table = document.add_table(rows=1, cols=2)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'日期'
        hdr_cells[1].text = u'內容'
        for diary in diaries:
                row_cells = table.add_row().cells
                row_cells[0].text = str(timezone.localtime(diary.time).strftime("%b %d %Y %H:%M:%S"))
                row_cells[1].text = diary.memo

        # Prepare document for download
        # -----------------------------
        f = StringIO.StringIO()
        document.save(f)
        length = f.tell()
        f.seek(0)
        response = HttpResponse(
                f.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=' + docx_title
        response['Content-Length'] = length

        return response